"""将完整 Markdown 论文构建为可编辑、可渲染的数学建模竞赛 Word 文档。"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor
from lxml import etree


ROOT = Path(__file__).absolute().parents[2]
SOURCE = ROOT / "paper_full_draft.md"
OUTPUT = ROOT / "paper_full_draft.docx"
FIGURE_DIR = ROOT / "figures" / "paper"
MML2OMML = Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")
EXTRA_DEPS = Path(r"C:\Users\haha\.codex\tmp\docx-deps")
if str(EXTRA_DEPS) not in sys.path:
    sys.path.insert(0, str(EXTRA_DEPS))

from latex2mathml.converter import convert as latex_to_mathml  # noqa: E402


INK = "1F2933"
ACCENT = "2F5D62"
ACCENT_DARK = "214448"
MUTED = "66737A"
TABLE_HEADER = "DCE9E7"
TABLE_ALT = "F5F8F8"
TABLE_BORDER = "A8B9B7"
CODE_FILL = "F3F5F6"
CONTENT_WIDTH_DXA = 9072
TABLE_INDENT_DXA = 120


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: float | None = None,
                 bold: bool | None = None, italic: bool | None = None, color: str | None = None) -> None:
    """显式设置中西文字体，避免 Word 与 LibreOffice 字体回退不一致。"""
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), TABLE_BORDER)


def set_table_geometry(table, widths: list[int]) -> None:
    """将表宽、缩进、网格和单元格宽度统一为固定 DXA 几何。"""
    tbl_pr = table._tbl.tblPr
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def text_weight(text: str) -> float:
    return sum(2.0 if "\u4e00" <= char <= "\u9fff" else 1.0 for char in text)


def compute_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    weights = []
    for col in range(cols):
        longest = max(text_weight(row[col]) for row in rows)
        weights.append(max(5.0, min(longest, 34.0)))
    if cols >= 3:
        for col in range(cols):
            numeric = sum(bool(re.fullmatch(r"[\d.,%+\- /()—-]+", row[col].strip())) for row in rows[1:])
            if numeric >= max(1, len(rows[1:]) // 2):
                weights[col] *= 0.72
    total = sum(weights)
    widths = [max(900, int(CONTENT_WIDTH_DXA * weight / total)) for weight in weights]
    scale = CONTENT_WIDTH_DXA / sum(widths)
    widths = [int(width * scale) for width in widths]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for item in (begin, instr, separate, text, end):
        run._r.append(item)


def mathml_to_omml(latex: str):
    latex = re.sub(r"\\tag\{[^{}]+\}", "", latex).strip()
    latex = latex.replace(r"\qquad", r"\quad\quad")
    mathml = latex_to_mathml(latex)
    tree = etree.fromstring(mathml.encode("utf-8"))
    transform = etree.XSLT(etree.parse(str(MML2OMML)))
    omml = transform(tree)
    return parse_xml(etree.tostring(omml.getroot()))


def append_inline_content(paragraph, text: str, size: float = 10.5, bold: bool = False) -> None:
    """解析粗体、斜体、行内代码和行内公式，保留 Word 可编辑结构。"""
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`|\$[^$]+?\$|\*[^*]+?\*)")
    position = 0
    for match in token_re.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size, bold=bold, color=INK)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=INK)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, east_asia="等线", latin="Consolas", size=max(8.5, size - 0.5), color=ACCENT_DARK)
            run.font.highlight_color = None
        elif token.startswith("$"):
            try:
                paragraph._p.append(mathml_to_omml(token[1:-1]))
            except Exception:
                run = paragraph.add_run(token[1:-1])
                set_run_font(run, east_asia="Cambria Math", latin="Cambria Math", size=size, italic=True, color=INK)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, italic=True, color=INK)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, bold=bold, color=INK)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.33

    heading_specs = {
        "Heading 1": (15, 15, 8, "黑体", ACCENT_DARK),
        "Heading 2": (12.5, 11, 5, "黑体", ACCENT),
        "Heading 3": (11, 8, 4, "黑体", ACCENT_DARK),
    }
    for name, (size, before, after, font, color) in heading_specs.items():
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(7)
    caption.paragraph_format.line_spacing = 1.15

    core = doc.core_properties
    core.title = "生产企业原材料的订购与运输：基于词典序 MILP/LP 与独立验证"
    core.subject = "2021-C 数学建模竞赛论文完整初稿"
    core.author = ""
    core.last_modified_by = ""
    core.comments = "由既有模型、结果和验证工件生成；未重新求解。"

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("生产企业原材料的订购与运输")
    set_run_font(run, east_asia="等线", latin="Arial", size=8.5, color=MUTED)
    p_pr = hp._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "C8D4D2")
    border.append(bottom)
    p_pr.append(border)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    add_field(fp, " PAGE ", "1")


def add_title(doc: Document, title: str) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(32)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(title)
    set_run_font(run, east_asia="黑体", latin="Arial", size=20, bold=True, color=ACCENT_DARK)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.first_line_indent = Pt(0)
    sub.paragraph_format.space_after = Pt(24)
    run = sub.add_run("数学建模竞赛论文完整初稿")
    set_run_font(run, east_asia="等线", latin="Arial", size=11, color=MUTED)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph("目录", style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(12)
    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        match = re.match(r"^(#{2,3})\s+(.+)$", line.strip())
        if not match:
            continue
        level = len(match.group(1)) - 1
        text = match.group(2).strip()
        toc = doc.add_paragraph()
        toc.paragraph_format.first_line_indent = Pt(0)
        toc.paragraph_format.left_indent = Cm(0 if level == 1 else 0.8)
        toc.paragraph_format.space_before = Pt(0)
        toc.paragraph_format.space_after = Pt(2.5 if level == 1 else 1.5)
        toc.paragraph_format.line_spacing = 1.05
        run = toc.add_run(text)
        set_run_font(
            run,
            east_asia="黑体" if level == 1 else "宋体",
            latin="Arial" if level == 1 else "Times New Roman",
            size=9.5 if level == 1 else 8.5,
            bold=(level == 1),
            color=ACCENT_DARK if level == 1 else INK,
        )
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_display_equation(doc: Document, latex: str) -> None:
    tag_match = re.search(r"\\tag\{([^{}]+)\}", latex)
    tag = tag_match.group(1) if tag_match else None
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_together = True
    tabs = paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(8.0), WD_TAB_ALIGNMENT.CENTER)
    tabs.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.add_run("\t")
    try:
        paragraph._p.append(mathml_to_omml(latex))
    except Exception:
        run = paragraph.add_run(re.sub(r"\\tag\{[^{}]+\}", "", latex).strip())
        set_run_font(run, east_asia="Cambria Math", latin="Cambria Math", size=10.5, italic=True, color=INK)
    if tag:
        run = paragraph.add_run(f"\t({tag})")
        set_run_font(run, size=10, color=INK)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    widths = compute_widths(rows)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_geometry(table, widths)
    set_table_borders(table)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row_index, (word_row, data_row) in enumerate(zip(table.rows, rows)):
        for col_index, (cell, value) in enumerate(zip(word_row.cells, data_row)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, TABLE_HEADER)
            elif row_index % 2 == 0:
                set_cell_shading(cell, TABLE_ALT)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(1.5)
            paragraph.paragraph_format.space_after = Pt(1.5)
            paragraph.paragraph_format.line_spacing = 1.1
            numeric = bool(re.fullmatch(r"[\d.,%+\- /()—-]+", value.strip()))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 or numeric else WD_ALIGN_PARAGRAPH.LEFT
            append_inline_content(paragraph, value, size=8.5, bold=(row_index == 0))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def add_image(doc: Document, path_text: str, alt: str) -> None:
    path = (ROOT / path_text).resolve()
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Cm(15.8))
    shape._inline.docPr.set("descr", alt or path.stem)


def add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Cm(0.55)
    paragraph.paragraph_format.right_indent = Cm(0.35)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.05
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CODE_FILL)
    p_pr.append(shd)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, east_asia="等线", latin="Consolas", size=8.5, color=INK)


def parse_table_line(line: str) -> list[str]:
    text = line.strip()
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|"):
        text = text[:-1]

    cells: list[str] = []
    buffer: list[str] = []
    in_math = False
    for index, char in enumerate(text):
        if char == "$" and (index == 0 or text[index - 1] != "\\"):
            in_math = not in_math
        if char == "|" and not in_math:
            cells.append("".join(buffer).strip())
            buffer = []
        else:
            buffer.append(char)
    cells.append("".join(buffer).strip())
    return cells


def build_document() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    title = lines[0].removeprefix("# ").strip()
    add_title(doc, title)

    index = 1
    paragraph_buffer: list[str] = []
    toc_added = False

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(item.strip() for item in paragraph_buffer).strip()
        paragraph_buffer = []
        if not text:
            return
        is_figure_caption = bool(re.match(r"\*\*图\s+\d+\s*\|", text))
        is_table_caption = bool(re.match(r"\*\*表\s+\d+\b", text))
        is_keywords = text.startswith("**关键词：**")
        is_reference = bool(re.match(r"\[\d+\]", text))
        if is_figure_caption or is_table_caption:
            paragraph = doc.add_paragraph(style="Caption")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.keep_with_next = is_table_caption
            append_inline_content(paragraph, text, size=9)
        elif is_keywords:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(8)
            append_inline_content(paragraph, text, size=10, bold=False)
        elif is_reference:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(-18)
            paragraph.paragraph_format.left_indent = Pt(18)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.15
            append_inline_content(paragraph, text, size=9)
        else:
            paragraph = doc.add_paragraph()
            append_inline_content(paragraph, text)

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("$$"):
            flush_paragraph()
            equation_lines = []
            index += 1
            while index < len(lines) and lines[index].strip() != "$$":
                equation_lines.append(lines[index])
                index += 1
            add_display_equation(doc, "\n".join(equation_lines).strip())
        elif stripped.startswith("```"):
            flush_paragraph()
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(doc, code_lines)
        elif re.match(r"^#{2,4}\s+", stripped):
            flush_paragraph()
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            if not toc_added and text.startswith("1 "):
                add_toc(doc)
                toc_added = True
            paragraph = doc.add_paragraph(text, style=f"Heading {min(level - 1, 3)}")
            if text in {"摘要", "16 参考文献", "17 附录"}:
                paragraph.paragraph_format.page_break_before = text != "摘要"
        elif stripped.startswith("!["):
            flush_paragraph()
            match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
            if match:
                add_image(doc, match.group(2), match.group(1))
        elif stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[index + 1].strip()):
            flush_paragraph()
            table_rows = [parse_table_line(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(parse_table_line(lines[index]))
                index += 1
            index -= 1
            add_markdown_table(doc, table_rows)
        elif re.match(r"^[-*]\s+", stripped):
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_after = Pt(3)
            append_inline_content(paragraph, re.sub(r"^[-*]\s+", "", stripped))
        elif re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_after = Pt(3)
            append_inline_content(paragraph, re.sub(r"^\d+\.\s+", "", stripped))
        elif not stripped:
            flush_paragraph()
        else:
            paragraph_buffer.append(line)
        index += 1

    flush_paragraph()

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size is None and paragraph.style.name == "Normal":
                set_run_font(run, size=10.5, color=INK)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
